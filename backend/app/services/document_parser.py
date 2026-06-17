"""Document parser service.

Extracts plain text content from various document formats (TXT, MD, PDF, DOCX).
"""

import logging
from io import BytesIO
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class DocumentParserError(Exception):
    """Raised when document parsing fails."""
    pass

class DocumentParser:
    @staticmethod
    def parse_txt(content: bytes) -> str:
        """Parse plain text content."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except Exception as e:
                raise DocumentParserError(f"Failed to decode text file: {e}")

    @staticmethod
    def parse_md(content: bytes) -> str:
        """Parse Markdown content (currently treated as plain text)."""
        return DocumentParser.parse_txt(content)

    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """Extract text from a PDF file using pypdf."""
        try:
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.exception("PDF parsing error")
            raise DocumentParserError(f"Failed to parse PDF document: {e}")

    @staticmethod
    def parse_docx(content: bytes) -> str:
        """Extract text from a DOCX file using python-docx."""
        try:
            doc = Document(BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)
            # Handle tables too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except Exception as e:
            logger.exception("DOCX parsing error")
            raise DocumentParserError(f"Failed to parse DOCX document: {e}")

    @staticmethod
    def parse_document(filename: str, content: bytes) -> str:
        """Parse document based on filename extension."""
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return DocumentParser.parse_pdf(content)
        elif ext in ("docx", "doc"):
            return DocumentParser.parse_docx(content)
        elif ext == "md":
            return DocumentParser.parse_md(content)
        elif ext in ("txt", "text", "csv"):
            return DocumentParser.parse_txt(content)
        else:
            # Fall back to text decoding
            try:
                return DocumentParser.parse_txt(content)
            except Exception:
                raise DocumentParserError(f"Unsupported document format: {ext}")
